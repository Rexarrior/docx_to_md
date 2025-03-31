import os
from docx import Document

def analyze_document(file_path):
    """Analyze the structure of a docx file."""
    doc = Document(file_path)
    
    print(f"Document: {os.path.basename(file_path)}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    
    print("\nHeading structure:")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') or p.text.strip() and p.runs and p.runs[0].bold:
            print(f"  {i}: Style: {p.style.name} - Text: {p.text[:50]}")
    
    print("\nSample paragraphs with formatting:")
    for i, p in enumerate(doc.paragraphs[:30]):
        if p.text.strip():
            print(f"  {i}: {p.text[:50]}")
            for j, run in enumerate(p.runs):
                format_info = []
                if run.bold: format_info.append("bold")
                if run.italic: format_info.append("italic")
                if run.underline: format_info.append("underline")
                if run.font.strike: format_info.append("strikethrough")
                
                if format_info or run.text.strip():
                    print(f"    Run {j}: {run.text[:30]} - Format: {', '.join(format_info)}")
    
    print("\nTable information:")
    for i, table in enumerate(doc.tables):
        print(f"  Table {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Print first row of each table
        if table.rows:
            print(f"    First row content:")
            for j, cell in enumerate(table.rows[0].cells):
                print(f"      Cell {j+1}: {cell.text[:30]}")
    
    print("\nDocument images:")
    image_parts = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_parts.append(rel.target_ref)
    print(f"  Found {len(image_parts)} images")
    for i, img in enumerate(image_parts):
        print(f"    Image {i+1}: {img}")
    
if __name__ == "__main__":
    analyze_document("docs/example_min.docx") 