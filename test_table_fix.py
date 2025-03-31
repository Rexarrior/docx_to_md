import os
from src.parser.docx_parser import DocxParser
from docx import Document as DocxDocument

def main():
    # Path to the example file
    example_file = "docs/example_min.docx"
    
    # Ensure the file exists
    if not os.path.exists(example_file):
        print(f"Example file not found: {example_file}")
        return
    
    print("Testing table detection in document flow...")
    
    # Load the docx file directly for comparison
    docx = DocxDocument(example_file)
    
    # Get all document elements
    body = docx._body._body
    
    # Count tables directly in document body flow
    tables_in_flow = []
    paragraphs_in_flow = []
    
    # Track positions for analysis
    element_count = 0
    for element in body.iterchildren():
        element_count += 1
        
        if element.tag.endswith('tbl'):
            tables_in_flow.append(element_count)
            
        if element.tag.endswith('p'):
            paragraphs_in_flow.append(element_count)
    
    print(f"Total elements in document: {element_count}")
    print(f"Tables found in document flow at positions: {tables_in_flow}")
    print(f"Total paragraphs in document flow: {len(paragraphs_in_flow)}")
    
    # Parse the document with our parser
    parser = DocxParser()
    document = parser.parse_document(example_file)
    
    # Count tables in parsed document
    table_paragraphs = []
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.table:
            table_paragraphs.append(i)
    
    print(f"\nOUR PARSER RESULTS:")
    print(f"Tables found by our parser: {len(table_paragraphs)}")
    print(f"Table positions in paragraphs: {table_paragraphs}")
    
    # Verify the tables are in the correct position
    if len(tables_in_flow) == len(table_paragraphs):
        print("\nSUCCESS: Our parser found the same number of tables as in the document flow.")
    else:
        print(f"\nERROR: Mismatch in table count. Document has {len(tables_in_flow)}, parser found {len(table_paragraphs)}.")
    
    # Look at content around the tables
    print("\nContext around tables:")
    for table_pos in table_paragraphs:
        print(f"\nTable at position {table_pos}:")
        
        # Show content of previous paragraphs for context
        start_idx = max(0, table_pos - 2)
        for i in range(start_idx, table_pos):
            print(f"  Para {i}: {document.paragraphs[i].text[:50]}...")
            
        # Show table info
        table = document.paragraphs[table_pos].table
        print(f"  TABLE: {table.num_rows} rows, {table.num_cols} columns")
        if table.num_rows > 0 and table.num_cols > 0:
            print(f"  First row: {table.rows[0]}")
            
        # Show content of next paragraphs for context
        end_idx = min(len(document.paragraphs) - 1, table_pos + 2)
        for i in range(table_pos + 1, end_idx + 1):
            print(f"  Para {i}: {document.paragraphs[i].text[:50]}...")

if __name__ == "__main__":
    main() 