"""
Special test script to debug table detection in the docx parser.
"""
import os
from docx import Document as DocxDocument
from src.parser.docx_parser import DocxParser

def main():
    # Path to the example document
    example_docx = os.path.join('docs', 'example_min.docx')
    
    if not os.path.exists(example_docx):
        print(f"Error: Example file {example_docx} not found")
        return
    
    # Load the docx directly to inspect table content
    print("DIRECT DOCX INSPECTION:")
    docx = DocxDocument(example_docx)
    print(f"Tables found in docx: {len(docx.tables)}")
    print(f"Paragraphs found in docx: {len(docx.paragraphs)}")
    
    # If tables exist, print some information about them
    for i, table in enumerate(docx.tables):
        print(f"\nTable {i+1}:")
        print(f"  Rows: {len(table.rows)}")
        print(f"  Columns: {len(table.columns)}")
        
        # Print first row if available
        if table.rows:
            print("  First row content:")
            for j, cell in enumerate(table.rows[0].cells):
                print(f"    Cell {j+1}: '{cell.text}'")
    
    # Try to locate tables within the document flow
    print("\nLOCATING TABLES IN DOCUMENT FLOW:")
    
    # Search for paragraphs that might be before or after tables
    table_indicators = ['Воркшопы', 'Дедлайны', 'Таблиц', 'Table']
    
    for i, para in enumerate(docx.paragraphs):
        # Look for paragraphs that might relate to tables
        if any(indicator in para.text for indicator in table_indicators):
            print(f"\nPossible table indicator in paragraph {i+1}:")
            print(f"  Text: '{para.text}'")
            
            # Check the surrounding paragraphs
            start = max(0, i-1)
            end = min(len(docx.paragraphs), i+3)
            print(f"  Context paragraphs ({start+1}-{end}):")
            for j in range(start, end):
                print(f"    Para {j+1}: '{docx.paragraphs[j].text}'")
    
    # Look for block-level structures in the document
    print("\nDOCUMENT STRUCTURE ANALYSIS:")
    try:
        # Sometimes we can inspect the document body
        body = docx._body._body
        for i, el in enumerate(body):
            tag = el.tag.split('}')[-1]
            if tag == 'tbl':
                print(f"  Block {i+1}: TABLE")
            elif tag == 'p':
                # Get paragraph text if available
                text = ""
                for child in el:
                    if child.tag.endswith('r'):  # Text run
                        for t in child.itertext():
                            text += t
                print(f"  Block {i+1}: PARAGRAPH '{text[:30]}...'")
            else:
                print(f"  Block {i+1}: {tag}")
    except Exception as e:
        print(f"  Error analyzing document structure: {str(e)}")
    
    # Parse with our parser and see what it finds
    print("\nOUR PARSER RESULTS:")
    parser = DocxParser()
    document = parser.parse_document(example_docx)
    
    # Count tables in our parsed document
    table_count = 0
    table_paragraphs = []
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.has_table():
            table_count += 1
            table_paragraphs.append(i+1)
    
    print(f"Tables found by our parser: {table_count}")
    print(f"Table paragraphs: {table_paragraphs}")
    
    if table_paragraphs:
        # Print paragraphs around the first table
        table_idx = table_paragraphs[0] - 1
        start = max(0, table_idx - 3)
        end = min(len(document.paragraphs), table_idx + 2)
        
        print("\nContext around first table:")
        for i in range(start, end):
            text = document.paragraphs[i].text
            text_preview = text[:30] + "..." if len(text) > 30 else text
            print(f"  Paragraph {i+1}: '{text_preview}'")
            if document.paragraphs[i].has_table():
                print(f"    HAS TABLE: {document.paragraphs[i].table.num_rows} rows")

if __name__ == "__main__":
    main() 