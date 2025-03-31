from typing import Dict, List, Any, Optional, Tuple
import os
from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentClass
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
# from docx.oxml.text.run import CT_R
# from docx.oxml.xmlchemy import OxmlElement

from ..models import Document, Paragraph, Image, Table

class DocxParser:
    """
    Parser for DOCX files.
    Converts DOCX elements into the document model structure.
    """
    
    def __init__(self):
        self.image_counter: int = 1  # Counter for generating image filenames
        self.document: Optional[Document] = None
        self.docx: Optional[DocxDocumentClass] = None
        
    def parse_document(self, docx_path: str) -> Document:
        """
        Parse a DOCX file and return a Document object.
        
        Args:
            docx_path (str): Path to the DOCX file
            
        Returns:
            Document: Parsed document object
        """
        # Create a new document
        self.document = Document()
        self.document.filename = os.path.basename(docx_path)
        
        # Load the docx file
        self.docx = DocxDocument(docx_path)
        
        # Process document body elements in order
        self._process_document_elements()
        
        return self.document
        
    def _process_document_elements(self):
        """
        Process all document elements (paragraphs and tables) in the order they appear.
        """
        # Get all block elements in document body
        body = self.docx._body._body
        
        # Process elements in document order
        for element in body.iterchildren():
            if element.tag.endswith('p'):
                # It's a paragraph
                paragraph = DocxParagraph(element, self.docx)
                if paragraph.text.strip() or self._has_image(paragraph):
                    parsed_paragraph = self._parse_paragraph(paragraph)
                    self.document.add_paragraph(parsed_paragraph)
            elif element.tag.endswith('tbl'):
                # It's a table
                table = DocxTable(element, self.docx)
                parsed_table = self._parse_table(table)
                
                # Create a paragraph for the table
                table_paragraph = Paragraph()
                table_paragraph.table = parsed_table
                
                self.document.add_paragraph(table_paragraph)
        
    def _parse_paragraph(self, paragraph: DocxParagraph) -> Paragraph:
        """
        Parse a paragraph element from docx
        
        Args:
            paragraph: A docx paragraph object
            
        Returns:
            Paragraph: The parsed Paragraph object
        """
        model_paragraph = Paragraph()
        
        # Extract text
        model_paragraph.text = paragraph.text
        
        # Detect if it's a heading
        model_paragraph.heading_level = self._detect_heading_level(paragraph)
        if model_paragraph.heading_level > 0:
            model_paragraph.heading = paragraph.text
        
        # Extract formatting
        model_paragraph.runs = self._extract_text_formatting(paragraph)
        
        # Extract image if present
        model_paragraph.image = self._extract_image(paragraph)
        
        return model_paragraph
        
    def _extract_image(self, paragraph: DocxParagraph) -> Optional[Image]:
        """
        Extract image from a paragraph
        
        Args:
            paragraph: A docx paragraph that might contain an image
            
        Returns:
            Optional[Image]: The extracted image or None
        """
        # Check if paragraph contains an image using xpath to find inline shapes
        image_elements = paragraph._element.xpath('.//a:blip')
        if not image_elements:
            return None
        
        # For simplicity, we'll just take the first image
        for element in paragraph._element.xpath('.//a:blip'):
            # Get the relationship ID
            rId = element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId:
                try:
                    # Get the image part
                    image_part = paragraph.part.related_parts[rId]
                    
                    # Create an Image object
                    image = Image()
                    image.file_path = image_part.partname
                    image.content = image_part.blob
                    
                    # Set alt text if available
                    try:
                        drawing_element = element.getparent().getparent().getparent()
                        doc_pr = drawing_element.xpath('.//wp:docPr', namespaces=drawing_element.nsmap)
                        if doc_pr:
                            image.alt_text = doc_pr[0].get('descr', '')
                        else:
                            image.alt_text = ""
                    except:
                        image.alt_text = ""
                    
                    # Generate a new filename
                    image.new_file_name = f"document.{self.image_counter:03d}.png"
                    self.image_counter += 1
                    
                    return image
                except:
                    # If we can't extract the image, continue to the next one
                    continue
        
        return None
        
    def _parse_table(self, table: DocxTable) -> Table:
        """
        Parse a table element from docx
        
        Args:
            table: A docx table object
            
        Returns:
            Table: The parsed Table object
        """
        model_table = Table()
        
        # Process rows
        for row in table.rows:
            model_row = []
            
            # Process cells
            for cell in row.cells:
                model_row.append(cell.text)
            
            model_table.add_row(model_row)
        
        # The first row is typically a header in markdown
        model_table.header = True
        
        # Set default alignments (left for all columns)
        for i in range(model_table.num_cols):
            model_table.set_column_alignment(i, 'left')
        
        return model_table
    
    def _get_font_size(self, run) -> Optional[int]:
        """
        Get the font size from a run
        
        Args:
            run: A docx run object
            
        Returns:
            Optional[int]: The font size in points or None if not found
        """
        # Try different ways to get font size (python-docx has different apis in different versions)
        if hasattr(run, 'font') and hasattr(run.font, 'size'):
            if run.font.size is not None:
                # Convert to points if needed (size can be in half-points or other units)
                try:
                    # Size might be a length object that needs conversion
                    if hasattr(run.font.size, 'pt'):
                        return int(run.font.size.pt)
                    # Size might be a raw value
                    return int(run.font.size) // 2  # Convert half-points to points
                except (ValueError, TypeError):
                    # If conversion fails, try to get raw value
                    return None
        
        # Try to get size from the XML directly if python-docx API didn't work
        try:
            if hasattr(run, '_element') and hasattr(run._element, 'xpath'):
                size_elements = run._element.xpath('.//w:sz')
                if size_elements:
                    size_val = size_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if size_val:
                        return int(size_val) // 2  # Convert half-points to points
        except:
            pass
        
        return None

    def _detect_heading_level(self, paragraph: DocxParagraph) -> int:
        """
        Detect the heading level of a paragraph
        
        Args:
            paragraph: A docx paragraph
            
        Returns:
            int: The heading level (0 for normal paragraphs)
        """
        # Check if it's a heading by style name - this is the most reliable method
        if hasattr(paragraph, 'style') and paragraph.style and paragraph.style.name.startswith('Heading'):
            try:
                # Extract heading level from style name (e.g., 'Heading 1' -> 1)
                return int(paragraph.style.name.split()[-1])
            except:
                # Default to level 1 if we can't extract the level
                return 1
                
        # Check for other heading styles that might not start with "Heading"
        # Some documents use Title, Subtitle, etc.
        if hasattr(paragraph, 'style') and paragraph.style:
            style_name = paragraph.style.name.lower()
            if 'title' in style_name:
                return 1
            if 'subtitle' in style_name:
                return 2
        
        # If no heading style is found, try to determine level from formatting
        if paragraph.runs and paragraph.text.strip():
            # Get the font size of the first run if available
            font_size = None
            if len(paragraph.runs) > 0:
                font_size = self._get_font_size(paragraph.runs[0])
            
            # Check if the entire paragraph is bold
            all_text = "".join(run.text for run in paragraph.runs)
            all_bold = True
            
            for run in paragraph.runs:
                if run.text.strip() and not run.bold:
                    all_bold = False
                    break
            
            # If it's all bold, determine level by additional clues
            if all_bold and all_text.strip():
                # Check if there's a font size we can use to determine level
                if font_size:
                    # Font sizes often decrease as heading levels increase
                    # These thresholds are estimates and may need adjustment
                    if font_size >= 20:  # Very large text
                        return 1
                    elif font_size >= 16:
                        return 2
                    elif font_size >= 14:
                        return 3
                    else:
                        return 4
                
                # Check for other formatting that might indicate level
                all_caps = all(c.isupper() for c in all_text if c.isalpha())
                if all_caps:
                    return 1  # ALL CAPS text is often a top-level heading
                
                # Look at length as a heuristic - shorter headings tend to be higher level
                if len(all_text.strip()) <= 20:
                    return 2
                else:
                    return 3
            
            # If the first run is bold and contains the entire text
            first_run = paragraph.runs[0]
            if first_run.bold and first_run.text.strip() and first_run.text.strip() == paragraph.text.strip():
                # Apply the same font size logic for single-run bold paragraphs
                if font_size:
                    if font_size >= 20:
                        return 1
                    elif font_size >= 16:
                        return 2
                    elif font_size >= 14:
                        return 3
                    else:
                        return 4
                
                # If no font size info, use text length as a basic heuristic
                if len(first_run.text.strip()) <= 20:
                    return 2
                else:
                    return 3
        
        return 0  # Not a heading
    
    def _extract_text_formatting(self, paragraph: DocxParagraph) -> List[Dict[str, Any]]:
        """
        Extract text runs with their formatting
        
        Args:
            paragraph: A docx paragraph
            
        Returns:
            List[Dict]: List of runs with text and formatting info
        """
        formatted_runs = []
        
        for run in paragraph.runs:
            if not run.text:
                continue
                
            # Extract formatting attributes
            formatting = {
                "text": run.text,
                "bold": bool(run.bold),
                "italic": bool(run.italic),
                "underline": bool(run.underline),
                "strike": bool(run.font.strike)
            }
            
            formatted_runs.append(formatting)
            
        return formatted_runs
    
    def _has_image(self, paragraph: DocxParagraph) -> bool:
        """
        Check if a paragraph contains an image
        
        Args:
            paragraph: A docx paragraph
            
        Returns:
            bool: True if the paragraph contains an image
        """
        return len(paragraph._element.xpath('.//a:blip')) > 0 